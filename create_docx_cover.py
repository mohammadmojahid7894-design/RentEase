import os
import sys
import subprocess

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

def add_centered_run(paragraph, text, pt_size=12, bold=False, color=None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(pt_size)
    run.font.name = 'Montserrat'
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

doc = Document()

# Set all document margins to 1 inch
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Header
p_header = doc.add_paragraph()
add_centered_run(p_header, "DHARANIDHAR UNIVERSITY, KEONJHAR\n", 22, bold=True, color=(17, 17, 17))
add_centered_run(p_header, "(Erstwhile Dharanidhar Auto College, Keonjhar)\n", 11, bold=False, color=(85, 85, 85))
add_centered_run(p_header, "Maharaja Sriram Chandra Bhanja Deo University", 11, bold=False, color=(85, 85, 85))

# Space
doc.add_paragraph()

# Logo
logo_path = "c:/Users/LENOVO/OneDrive/Desktop/documents/6th SEM PROJECT/ghar-ka-system (1)/university_logo.png"
if os.path.exists(logo_path):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run()
    r_logo.add_picture(logo_path, width=Inches(1.5))

# Space
doc.add_paragraph()

# Title Section
p_title = doc.add_paragraph()
add_centered_run(p_title, "PROJECT REPORT ON\n", 14, bold=False, color=(68, 68, 68))
run_title = add_centered_run(p_title, "RentEase", 24, bold=True, color=(212, 175, 55)) # Gold color
add_centered_run(p_title, " – Smart Rental\nProperty Management System", 20, bold=True, color=(17, 17, 17))

# Illustration
illus_path = "c:/Users/LENOVO/OneDrive/Desktop/documents/6th SEM PROJECT/ghar-ka-system (1)/rentease_illustration.png"
if os.path.exists(illus_path):
    p_illus = doc.add_paragraph()
    p_illus.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_illus.paragraph_format.space_before = Pt(12)
    p_illus.paragraph_format.space_after = Pt(12)
    r_illus = p_illus.add_run()
    r_illus.add_picture(illus_path, width=Inches(3.5))

# Details
p_details = doc.add_paragraph()
add_centered_run(p_details, "Submitted in Partial Fulfillment for Award of\n", 12, color=(68, 68, 68))
add_centered_run(p_details, "Degree in B.Sc in Computer Science\n", 14, bold=True, color=(17, 17, 17))
add_centered_run(p_details, "(BATCH 2022 – 2025)", 12, bold=True, color=(212, 175, 55))

doc.add_paragraph() # space

# Student & Guide using Table
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
table.autofit = False
table.allow_autofit = False

# Adjust column widths
for cell in table.columns[0].cells:
    cell.width = Inches(3.0)
for cell in table.columns[1].cells:
    cell.width = Inches(3.0)

# Student Info
cell_student = table.cell(0, 0)
p_s = cell_student.paragraphs[0]
p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_centered_run(p_s, "Submitted By:\n", 12, bold=True, color=(102, 102, 102))
add_centered_run(p_s, "BISWAJIT SAHOO\n", 14, bold=True, color=(17, 17, 17))
add_centered_run(p_s, "Roll No: 2322D011\nRegd No: 17346/22", 11, color=(85, 85, 85))

# Guide Info
cell_guide = table.cell(0, 1)
p_g = cell_guide.paragraphs[0]
p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_centered_run(p_g, "Under the Esteemed Guidance of:\n", 12, bold=True, color=(102, 102, 102))
add_centered_run(p_g, "MR. SHUBHASHISH BEHERA\n", 14, bold=True, color=(17, 17, 17))
add_centered_run(p_g, "Department of Computer Science", 11, color=(85, 85, 85))

# Space at bottom
p_space = doc.add_paragraph()
p_space.paragraph_format.space_before = Pt(30)

# Footer Date
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_centered_run(p_footer, "_________________________\n", 14, color=(212, 175, 55))
add_centered_run(p_footer, "APRIL 2025", 14, bold=True, color=(17, 17, 17))

output_path = "c:/Users/LENOVO/OneDrive/Desktop/documents/6th SEM PROJECT/ghar-ka-system (1)/Front_Page.docx"
doc.save(output_path)
print(f"Word document saved successfully at {output_path}")
